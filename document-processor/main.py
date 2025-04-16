import os
import uuid
import tempfile
import logging
import asyncio
from typing import List, Dict, Any, Optional, Union, Tuple
from enum import Enum

import uvicorn
import httpx
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

# Document processing libraries
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
import re
import nltk
from nltk.tokenize import sent_tokenize

# Download NLTK resources if not already downloaded
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt_tab')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("document-processor")

# Document type enum
class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    TEXT = "txt"
    UNKNOWN = "unknown"

# Schema for document chunk
class DocumentChunk(BaseModel):
    chunk_id: str
    document_id: str
    text: str
    metadata: Dict[str, Any] = Field(default_factory=dict)

# Schema for processed document
class ProcessedDocument(BaseModel):
    document_id: str
    filename: str
    document_type: DocumentType
    chunk_count: int
    metadata: Dict[str, Any] = Field(default_factory=dict)
    status: str = "processed"

# Create FastAPI app
app = FastAPI(
    title="Document Processing Service",
    description="Service for processing documents, generating embeddings, and storing in Milvus",
    version="1.0.0",
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Adjust in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration from environment variables
CHROMA_HOST = os.getenv("CHROMA_DB_HOST", "chromadb")
CHROMA_PORT = int(os.getenv("CHROMA_DB_PORT", "8000"))
CHROMA_PERSIST_DIRECTORY = os.getenv("CHROMA_PERSIST_DIRECTORY", "./chroma_db")
OLLAMA_API_URL = os.getenv("OLLAMA_API_URL", "http://localhost:8000")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
COLLECTION_NAME = os.getenv("CHROMA_COLLECTION", "docs")
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "512"))  # Characters per chunk
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "128"))  # Overlap between chunks
MAX_CHUNKS_PER_DOC = int(os.getenv("MAX_CHUNKS_PER_DOC", "500"))  # Maximum chunks per document

# HTTP client
http_client = httpx.AsyncClient(timeout=60.0)

# Document processing functions
def detect_document_type(filename: str, content_type: str = None) -> DocumentType:
    """Detect document type from filename and content type."""
    extension = filename.lower().split('.')[-1]
    
    # First check based on file extension
    if extension == "pdf":
        return DocumentType.PDF
    elif extension in ["docx", "doc"]:
        return DocumentType.DOCX
    elif extension in ["txt", "text"]:
        return DocumentType.TEXT
    
    # If extension check doesn't provide a clear answer, check content type
    if content_type:
        if 'pdf' in content_type:
            return DocumentType.PDF
        elif 'word' in content_type or 'docx' in content_type or 'doc' in content_type:
            return DocumentType.DOCX
        elif 'text' in content_type or 'txt' in content_type or 'plain' in content_type:
            return DocumentType.TEXT
    
    # If still can't determine, try to open the file and check content
    try:
        with open(filename, 'rb') as f:
            # Check for PDF signature
            if f.read(4) == b'%PDF':
                return DocumentType.PDF
    except:
        pass
    
    # Default case - unknown
    return DocumentType.UNKNOWN

async def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    try:
        text = ""
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise

async def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise

async def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise

async def extract_text(file_path: str, document_type: DocumentType) -> str:
    """Extract text from a document based on its type."""
    if document_type == DocumentType.PDF:
        return await extract_text_from_pdf(file_path)
    elif document_type == DocumentType.DOCX:
        return await extract_text_from_docx(file_path)
    elif document_type == DocumentType.TEXT:
        return await extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported document type: {document_type}")

def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """Split text into chunks with overlap."""
    # Basic sentence tokenization to avoid cutting sentences
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed the chunk size, save the current chunk
        if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
            chunks.append(current_chunk)
            # Keep the overlap amount from the end of the current chunk
            current_chunk = current_chunk[-chunk_overlap:] if chunk_overlap < len(current_chunk) else current_chunk
        
        current_chunk += sentence + " "
    
    # Add the last chunk if it's not empty
    if current_chunk.strip():
        chunks.append(current_chunk)
    
    # Limit the number of chunks to prevent processing too many
    if len(chunks) > MAX_CHUNKS_PER_DOC:
        logger.warning(f"Document has {len(chunks)} chunks, limiting to {MAX_CHUNKS_PER_DOC}")
        chunks = chunks[:MAX_CHUNKS_PER_DOC]
    
    return chunks

async def setup_chroma_collection():
    """Setup ChromaDB connection and create collection if it doesn't exist."""
    try:
        # Get or create the ChromaDB client
        chroma_client = get_chroma_client()
        
        collection = chroma_client.get_or_create_collection(name=COLLECTION_NAME)

        return collection
    except Exception as e:
        logger.error(f"Failed to setup ChromaDB collection: {e}")
        raise

def get_chroma_client():
    """Get a ChromaDB client."""
    try:
        client = chromadb.HttpClient(host=f"{CHROMA_HOST}", port=CHROMA_PORT)
        client.heartbeat()

        return client
    except Exception as e:
        logger.error(f"Failed to connect to ChromaDB: {e}")
        raise



async def insert_chunks_to_chroma(chunks: List[DocumentChunk], embeddings: List[List[float]] = None):
    """Insert chunks and their embeddings into ChromaDB."""
    try:
        # Get the ChromaDB client and collection
        chroma_client = get_chroma_client()
        collection = chroma_client.get_collection(name=COLLECTION_NAME)
        
        # Prepare data for insertion
        ids = [chunk.chunk_id for chunk in chunks]
        documents = [chunk.text for chunk in chunks]
        metadatas = []
        
        # Log lengths for debugging
        logger.info(f"Inserting chunks to ChromaDB: ids={len(ids)}, documents={len(documents)}")
        
        # Process metadata to make it compatible with ChromaDB
        for chunk in chunks:
            # Make sure metadata is JSON serializable
            metadata = dict(chunk.metadata)
            metadata["document_id"] = chunk.document_id
            
            # Ensure all values are strings or simple types
            for key, value in list(metadata.items()):
                if isinstance(value, (dict, list)):
                    metadata[key] = str(value)
                elif not isinstance(value, (str, int, float, bool, type(None))):
                    metadata[key] = str(value)
            
            metadatas.append(metadata)
        
        logger.info(f"Prepared metadata: {len(metadatas)} items")
        
        # Check that all arrays have the same length
        if not (len(ids) == len(documents) == len(metadatas)):
            raise ValueError(f"Inconsistent lengths: ids={len(ids)}, documents={len(documents)}, metadatas={len(metadatas)}")
        
        # Insert data in batches to avoid potential issues with large documents
        batch_size = 50
        for i in range(0, len(ids), batch_size):
            batch_end = min(i + batch_size, len(ids))
            batch_ids = ids[i:batch_end]
            batch_documents = documents[i:batch_end]
            batch_metadatas = metadatas[i:batch_end]
            
            logger.info(f"Inserting batch {i//batch_size + 1}: {len(batch_ids)} chunks")
            
            # Insert data
            collection.add(
                ids=batch_ids,
                documents=batch_documents,
                metadatas=batch_metadatas,
            )
        
        logger.info(f"Successfully inserted {len(chunks)} chunks into ChromaDB")
        return True
    except Exception as e:
        logger.error(f"Failed to insert chunks into ChromaDB: {str(e)}")
        # Log additional debug info
        if chunks:
            logger.error(f"First chunk ID: {chunks[0].chunk_id}, text length: {len(chunks[0].text)}")
            logger.error(f"Metadata keys: {chunks[0].metadata.keys() if chunks[0].metadata else 'None'}")
        raise

async def process_document(
    file_path: str,
    filename: str,
    document_type: DocumentType,
    metadata: Dict[str, Any]
) -> ProcessedDocument:
    """Process a document: extract text, chunk it, generate embeddings, and store in ChromaDB."""
    try:
        document_id = str(uuid.uuid4())
        
        # Extract text from document
        extracted_text = await extract_text(file_path, document_type)
        
        # Chunk the text - renamed variable to avoid conflict with function name
        text_chunks = chunk_text(extracted_text, CHUNK_SIZE, CHUNK_OVERLAP)
        
        document_chunks = []
        
        # Create chunk objects
        for i, chunk_content in enumerate(text_chunks):
            chunk_id = str(uuid.uuid4())
            
            # Create document chunk
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                document_id=document_id,
                text=chunk_content,
                metadata={
                    "chunk_index": i,
                    "filename": filename,
                    "document_type": document_type,
                    **metadata
                }
            )
            
            document_chunks.append(chunk)
        
        # Insert chunks into ChromaDB
        # ChromaDB will automatically generate embeddings using our custom embedding function
        await insert_chunks_to_chroma(document_chunks)
        
        return ProcessedDocument(
            document_id=document_id,
            filename=filename,
            document_type=document_type,
            chunk_count=len(document_chunks),
            metadata=metadata
        )
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise

# Endpoints
@app.on_event("startup")
async def startup_event():
    """Start up event handler."""
    # Setup ChromaDB collection
    await setup_chroma_collection()

@app.on_event("shutdown")
async def shutdown_event():
    """Shutdown event handler."""
    # Close HTTP client
    await http_client.aclose()
    # ChromaDB automatically closes connections

@app.get("/")
async def root():
    """Root endpoint returning service info."""
    return {
        "service": "Document Processing Service",
        "version": "1.0.0",
        "status": "running",
    }

@app.get("/health")
async def health():
    """Health check endpoint."""
    try:
        # Check ChromaDB connection
        chroma_client = get_chroma_client()
        chroma_health = COLLECTION_NAME in chroma_client.list_collections()
        
        # Check Ollama API
        ollama_response = await http_client.get(f"{OLLAMA_API_URL}/health")
        ollama_health = ollama_response.status_code == 200
        
        return {
            "status": "healthy" if chroma_health and ollama_health else "unhealthy",
            "chroma": "healthy" if chroma_health else "unhealthy",
            "ollama": "healthy" if ollama_health else "unhealthy",
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return {
            "status": "unhealthy",
            "error": str(e)
        }

async def process_file_background(
    file_path: str,
    filename: str,
    document_type: DocumentType,
    metadata: Dict[str, Any]
):
    """Background task to process a file."""
    try:
        await process_document(file_path, filename, document_type, metadata)
    except Exception as e:
        logger.error(f"Background processing failed: {e}")
    finally:
        # Remove the temporary file
        try:
            os.remove(file_path)
        except Exception as e:
            logger.error(f"Failed to remove temporary file: {e}")

@app.post("/process", response_model=Dict[str, Any])
async def process_file(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    metadata: str = Form("{}")
):
    """Endpoint to upload and process a document."""
    try:
        # Parse metadata
        try:
            import json
            metadata_dict = json.loads(metadata)
        except Exception:
            metadata_dict = {}
        
        # Detect document type using both filename and content type
        document_type = detect_document_type(file.filename, file.content_type)
        if document_type == DocumentType.UNKNOWN:
            # Log the content type for debugging
            logger.warning(f"Unsupported file type. Filename: {file.filename}, Content-Type: {file.content_type}")
            raise HTTPException(status_code=400, detail=f"Unsupported file type. Content-Type: {file.content_type}")
        
        # Save the uploaded file temporarily
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        try:
            contents = await file.read()
            with open(temp_file.name, 'wb') as f:
                f.write(contents)
        except Exception as e:
            os.unlink(temp_file.name)
            raise HTTPException(status_code=500, detail=f"Failed to save file: {e}")
        
        # Add file info to metadata
        metadata_dict["original_filename"] = file.filename
        metadata_dict["content_type"] = file.content_type or "unknown"
        
        # Add the processing task to background tasks
        background_tasks.add_task(
            process_file_background,
            temp_file.name,
            file.filename,
            document_type,
            metadata_dict
        )
        
        return {
            "status": "processing",
            "filename": file.filename,
            "document_type": document_type,
            "message": "File uploaded and processing started in the background"
        }
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/status")
async def service_status():
    """Get service status and statistics."""
    try:
        # Check ChromaDB collection
        chroma_client = get_chroma_client()
        collection = chroma_client.get_collection(name=COLLECTION_NAME)
        collection_count = collection.count()
        
        collection_info = {
            "document_count": collection_count
        }
        
        return {
            "status": "running",
            "collection": COLLECTION_NAME,
            "collection_info": collection_info,
            "embedding_model": EMBEDDING_MODEL,
            "chunk_size": CHUNK_SIZE,
            "chunk_overlap": CHUNK_OVERLAP
        }
    except Exception as e:
        logger.error(f"Status check failed: {e}")
        return {
            "status": "error",
            "error": str(e)
        }

if __name__ == "__main__":
    # Get port from environment or use default
    port = int(os.getenv("PORT", "8001"))
    
    # Run the server
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)